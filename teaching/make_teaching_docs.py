# -*- coding: utf-8 -*-
"""공동체 빌더스 교사용 지도안 + 학생 활동지 — 일러스트 디자인 v2"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/Users/yeolstudio/Claude/emotion-quest-rpg/teaching"
ASSETS = os.path.join(ROOT, "assets")
os.makedirs(ROOT, exist_ok=True)

# ---------- 팔레트 (게임 UI와 동일 톤) ----------
CORAL = RGBColor(0xE8, 0x5D, 0x3F)
CORAL_BG = "FDE9E1"
MINT = RGBColor(0x0F, 0x8F, 0x68)
MINT_BG = "DCF5EA"
SKY = RGBColor(0x22, 0x77, 0xB8)
SKY_BG = "DDEEFB"
PURPLE = RGBColor(0x6D, 0x4B, 0xC9)
PURPLE_BG = "EAE2FB"
GOLD = RGBColor(0xB4, 0x7B, 0x0B)
GOLD_BG = "FBF0D6"
INK = RGBColor(0x3A, 0x30, 0x28)
GRAY = RGBColor(0x8A, 0x80, 0x76)
CREAM = "FFF9F0"

PALETTE_CYCLE = [(CORAL, CORAL_BG), (MINT, MINT_BG), (SKY, SKY_BG), (PURPLE, PURPLE_BG)]


# ---------- 저수준 헬퍼 ----------
def _oxml_border_edge(color, sz):
    el = OxmlElement("w:tcBorders" + "")  # placeholder, unused
    return el


def set_cell_shading(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_borders(cell, color="E8DDCF", sz=6, edges=("top", "left", "bottom", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def prevent_row_split(table):
    """표/글상자가 페이지 경계에서 둘로 쪼개져 보이지 않도록 행 전체를 다음 페이지로 넘긴다"""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)


def set_row_height(row, cm_val):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(Cm(cm_val))))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


def set_page_border(section, color, sz=14, space=18):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


# ---------- 공통 빌더 ----------
def set_korean_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def new_doc(margins=(1.6, 1.6, 1.9, 1.9)):
    doc = Document()
    set_korean_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(margins[0])
    sec.bottom_margin = Cm(margins[1])
    sec.left_margin = Cm(margins[2])
    sec.right_margin = Cm(margins[3])
    set_page_border(sec, "F0C7B4", sz=10, space=16)
    return doc


def add_banner(doc, width_cm=17.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS, "banner_final.jpg"), width=Cm(width_cm))
    # 이미지 모서리를 살짝 정리하는 얇은 그라데이션 느낌의 바 (색 리본)
    return p


def add_title_ribbon(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4A186")
    set_cell_margins(cell, top=220, bottom=220, left=200, right=200)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xFF, 0xF3, 0xEC)
    prevent_row_split(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def section_chip(doc, icon_file, text, idx=0):
    """아이콘 + 색 배경의 둥근 느낌 섹션 헤더 칩"""
    accent, bg = PALETTE_CYCLE[idx % len(PALETTE_CYCLE)]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_table_borders(table)
    set_col_widths(table, [1.5, 15.05])
    icon_cell, text_cell = table.rows[0].cells
    set_cell_shading(icon_cell, bg)
    set_cell_shading(text_cell, bg)
    set_cell_margins(icon_cell, top=90, bottom=90, left=80, right=40)
    set_cell_margins(text_cell, top=90, bottom=90, left=40, right=140)
    icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if icon_file:
        ip = icon_cell.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(os.path.join(ASSETS, icon_file), width=Cm(1.2))
    tp = text_cell.paragraphs[0]
    run = tp.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = accent
    prevent_row_split(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return accent


def callout_box(doc, lines, accent, bg, icon_file=None, bold_first=False):
    """글상자: 색 테두리 + 옅은 배경의 강조 박스"""
    table = doc.add_table(rows=1, cols=1)
    no_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    set_cell_borders(cell, color="%02X%02X%02X" % (accent[0], accent[1], accent[2]), sz=10)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK
        if bold_first and p is cell.paragraphs[0]:
            run.bold = True
            run.font.color.rgb = accent
        p.paragraph_format.space_after = Pt(3)
    prevent_row_split(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def underline_run(paragraph, text, color=None, bold=False):
    run = paragraph.add_run(text)
    run.underline = True
    run.bold = bold
    run.font.color.rgb = color or CORAL
    return run


def body(doc, text_parts, size=10.5, space_after=5):
    """text_parts: 문자열이거나 (텍스트, {underline, bold, color}) 튜플의 리스트"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    if isinstance(text_parts, str):
        text_parts = [text_parts]
    for part in text_parts:
        if isinstance(part, tuple):
            text, opts = part
            run = p.add_run(text)
            run.font.size = Pt(opts.get("size", size))
            run.bold = opts.get("bold", False)
            run.underline = opts.get("underline", False)
            run.font.color.rgb = opts.get("color", INK)
        else:
            run = p.add_run(part)
            run.font.size = Pt(size)
    return p


def styled_grid(doc, rows, widths_cm, header_bg="F4E9DC", header_color=None, body_bg=None, font_size=9.5):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    set_col_widths(table, widths_cm)
    for r, row_vals in enumerate(rows):
        for c, text in enumerate(row_vals):
            cell = table.cell(r, c)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_borders(cell, color="EFE2D2", sz=5)
            if r == 0:
                set_cell_shading(cell, header_bg)
            elif body_bg:
                set_cell_shading(cell, body_bg if r % 2 == 0 else "FFFFFF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(font_size)
            if r == 0:
                run.bold = True
                run.font.color.rgb = header_color or INK
    prevent_row_split(table)
    return table


def divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("🌷 · 🌼 · 🌷 · 🌼 · 🌷")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD8, 0xA9, 0x8C)


# ==================================================================
# 1. 교사용 지도안
# ==================================================================
doc = new_doc()
add_banner(doc)
add_title_ribbon(doc, "사회정서교육(SEL) 수업 지도안", "공동체 빌더스 — 안개에 갇힌 마을을 함께 깨우기")

section_chip(doc, "icon_prep_final.png", "1. 수업 개요", 0)
styled_grid(doc, [
    ["항목", "내용"],
    ["대상", "중학교 1~3학년 (권장 학급 규모 20~35명)"],
    ["차시", "1차시 (45분)"],
    ["장소·기기", "컴퓨터실 또는 1인 1기기(태블릿·크롬북·노트북), 인터넷 연결"],
    ["접속 주소", "https://cleveranawim-source.github.io/Community_builders/"],
    ["SEL 영역", "자기인식(마음알기) · 대인관계(서로듣기·관계잇기) · 공동체(마을세우기)"],
], widths_cm=[3.4, 13.5], header_bg="F4E9DC")
doc.add_paragraph().paragraph_format.space_after = Pt(4)

section_chip(doc, "icon_goal_final.png", "2. 학습 목표", 1)
body(doc, [("① ", {"bold": True, "color": MINT}), "일상적인 학급 갈등 상황에서 공동체를 세우는 선택과 허무는 선택을 구별할 수 있다."])
body(doc, [("② ", {"bold": True, "color": MINT}), "경청·사과·합의·환대 등 관계 기술의 구체적 행동 방법을 말할 수 있다."])
body(doc, [("③ ", {"bold": True, "color": MINT}), "개인의 작은 실천이 공동체 전체를 밝힌다는 것을 ", ("우리 반 마을 밝기", {"underline": True, "bold": True, "color": MINT}), " 경험으로 설명할 수 있다."])

section_chip(doc, "icon_activity_final.png", "3. 수업 전 준비 (교사)", 2)
body(doc, "① 교사 컴퓨터에서 접속 주소를 열고, 시작 화면의 [교사용 화면 열기]를 클릭한다.")
body(doc, [("② ", {}), "[새 코드 만들기]를 눌러 반 코드를 만들고 칠판에 크게 적어 둔다. (예: ", ("3021", {"bold": True, "underline": True, "color": SKY}), ")"])
body(doc, "③ 교사 화면을 프로젝터/전자칠판에 띄워 두면 수업 내내 학생 진행 상황이 실시간으로 보인다.")
body(doc, "④ 학생 안내: 시작 화면에서 닉네임(실명 또는 별칭 규칙 정하기)과 반 코드를 입력하게 한다.")

section_chip(doc, "icon_activity_final.png", "4. 수업 흐름 (45분)", 2)
styled_grid(doc, [
    ["단계", "시간", "교수·학습 활동", "자료·유의점"],
    ["도입", "5분",
     "발문: \"우리 반이 '안개에 갇힌 마을'이라면, 그 안개는 무엇일까?\" (오해·침묵·소외 등)\n"
     "게임 소개: 안개 속에 숨은 16명의 친구를 찾아 고민을 함께 해결하면 마을이 깨어난다.\n"
     "접속: 닉네임 + 반 코드 입력",
     "칠판에 반 코드\n닉네임 규칙 안내"],
    ["전개1\n탐험", "20분",
     "자유 탐험: 빛장벽을 깨고, 안개 고치(물음표) 속 친구들을 발견해 미션을 해결한다.\n"
     "미션마다 두 선택지 중 하나를 고르고, 해설 카드를 반드시 읽게 한다.\n"
     "교사는 대시보드를 보며 진행이 느린 학생을 개별 지원한다.\n"
     "중간 환기(10분 경과): 우리 반 마을 밝기가 몇 %인지 함께 확인",
     "교사 화면 투사\n미션 6개 이상 권장\n(전체 완주 강요 금지)"],
    ["전개2\n나눔", "12분",
     "기기를 덮고 짝 활동: 활동지에 기억에 남는 미션 2가지와 그 이유를 기록\n"
     "짝과 나눔: 가장 공감된 해설 카드 서로 소개하기\n"
     "전체 나눔 2~3명: 우리 반에 실제로 필요한 약속 발표",
     "학생 활동지\n(1인 1장)"],
    ["정리", "8분",
     "교사 화면으로 '우리 반 마을 밝기'와 참여 현황을 함께 확인하며 격려\n"
     "연결 발문: \"게임 속 마을처럼, 우리 반의 안개는 무엇으로 걷힐까?\"\n"
     "우리 반 실천 약속 1가지 함께 정하기 (활동지 마지막 칸에 기록)\n"
     "완주한 학생은 인증서 저장 안내",
     "CSV 내보내기로\n수업 기록 저장"],
], widths_cm=[2.0, 1.5, 9.7, 3.2], header_bg="FDE9E1", header_color=CORAL)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

section_chip(doc, "icon_share_final.png", "5. 평가 관점 (관찰·활동지)", 3)
body(doc, "· 해설 카드의 핵심(예: 사과의 3단계, 다수결 전 걱정 듣기)을 자기 언어로 옮겨 적었는가")
body(doc, "· 게임 상황을 우리 반의 실제 상황과 연결 지어 말할 수 있는가")
body(doc, "· 짝 나눔에서 경청 태도(차례 지키기, 되묻기)를 보였는가")

section_chip(doc, "icon_tip_final.png", "6. 지도상 유의점", 0)
callout_box(doc, [
    "게임 순위 경쟁이 아니라 '반 전체 밝기'라는 공동 목표를 반복해서 강조합니다.",
    "'아쉬운 선택'을 골라도 정죄하지 않습니다 — 해설 카드를 읽고 다시 시도할 수 있음을 안내합니다.",
    "소리가 부담스러운 환경이면 화면 좌상단 🔊 버튼으로 음소거할 수 있습니다.",
    "진행 상황은 기기에 자동 저장되므로, 새로고침되어도 [이어서 하기]로 복구됩니다.",
    "문항을 학급 상황에 맞게 바꾸려면 quests-custom.json 파일을 수정합니다(안내서 참고).",
], CORAL, CORAL_BG, bold_first=False)

# ----- 새 페이지: 서울 SEL 성취기준 연계표 -----
doc.add_page_break()
add_banner(doc, width_cm=17.0)
section_chip(doc, "icon_goal_final.png", "7. 서울 사회정서교육 성취기준 연계 (중학교)", 1)

body(doc, [
    ("게임의 4역량이 서울 SEL 4영역과 거의 1:1로 대응합니다. ", {}),
    ("대인관계(02)와 공동체(03) 영역이 중심축", {"bold": True, "color": MINT}),
    ("이며, 자기(01)·마음건강(04)을 부분적으로 아우릅니다.", {}),
])

styled_grid(doc, [
    ["게임 역량", "서울 SEL 영역", "성격"],
    ["마음알기", "자기(01)", "감정 인식·회복탄력성"],
    ["서로듣기", "대인관계(02)", "공감 중심"],
    ["관계잇기", "대인관계(02)", "관계·갈등 해결"],
    ["마을세우기", "공동체(03)", "협력·공동체 가치"],
], widths_cm=[3.6, 5.0, 8.3], header_bg=MINT_BG, header_color=MINT)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

section_chip(doc, "icon_activity_final.png", "미션별 대표 성취기준", 2)
styled_grid(doc, [
    ["미션(마을 친구)", "대표 성취기준"],
    ["마음광장 복구(나리)", "[9사회정서01-01] 감정 인식·표현"],
    ["회복캠프(라온)", "[9사회정서01-05] 회복탄력성"],
    ["경청정원(민)", "[9사회정서02-01] 공감 · [02-04] 의사소통"],
    ["환영문(유나) / 칭찬무대(도윤)", "[9사회정서02-01] 공감"],
    ["관계다리 건설(레오)", "[9사회정서02-02] 관점 조율 · [02-05] 갈등"],
    ["모둠연결역(이안)", "[9사회정서02-02] 관점 · [03-02] 협력"],
    ["안전대화소(서우)", "[9사회정서02-03] 관계·경계 · [02-05] 갈등"],
    ["관계수리소(태오)", "[9사회정서02-05] 갈등 해결·관계 회복"],
    ["협동공방(소리) / 나눔창고(보미)", "[9사회정서03-02] 협력"],
    ["약속게시판(준)", "[9사회정서03-01] 공동체 가치 · [03-03] 책임"],
    ["합의광장(세아)", "[9사회정서03-03] 의사결정 · [03-04] 문제 해결"],
    ["평화쉼터(모루) / 도움신호 탑(하루)", "[9사회정서03-01] 공동체 · [04-02] 도움 자원"],
    ["또래도움소(가온)", "[9사회정서04-02] 도움 자원 · [02-01] 공감"],
], widths_cm=[6.5, 10.4], header_bg="FDE9E1", header_color=CORAL, font_size=9)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout_box(doc, [
    "게임 시스템도 기준에 닿습니다: '공감 빛구슬로 오해·침묵·소외 장벽 걷기'는 [9사회정서02-01](공감)의 은유적 실천, '우리 반 마을 밝기'(개인의 실천이 공동체를 밝힘)는 [9사회정서03-01·03-02](공동체 가치·협력)를 게임 메커니즘으로 구현한 것입니다.",
    "영역 균형: 이 게임은 '함께 세우기'가 강점이라 자기(01) 영역(감정 조절·스트레스)은 얇습니다. 자기인식·조절 활동과 짝을 지으면 4영역이 고르게 채워집니다.",
], SKY, SKY_BG, bold_first=False)

divider(doc)
doc.save(os.path.join(ROOT, "공동체빌더스_교사용_지도안.docx"))
print("지도안 저장 완료")


# ==================================================================
# 2. 학생 활동지
# ==================================================================
doc = new_doc()
add_banner(doc)
add_title_ribbon(doc, "공동체 빌더스", "마을을 깨운 나의 기록")

info_table = doc.add_table(rows=1, cols=1)
no_table_borders(info_table)
icell = info_table.rows[0].cells[0]
set_cell_shading(icell, "FFF9F0")
set_cell_borders(icell, color="F0C7B4", sz=8)
set_cell_margins(icell, top=100, bottom=100, left=200, right=200)
ip = icell.paragraphs[0]
ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = ip.add_run("학년 (        )    반 (        )    번호 (        )    이름 (                )")
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = INK
prevent_row_split(info_table)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

accent1 = section_chip(doc, "icon_activity_final.png", "1. 내가 만난 마을 친구들", 2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run("탐험하며 해결한 미션 중 기억에 남는 것을 골라 적어 보세요.")
run.italic = True
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
styled_grid(doc, [
    ["거점 이름", "친구 이름", "내가 고른 선택", "해설 카드에서 기억에 남는 문장"],
    ["", "", "", ""],
    ["", "", "", ""],
    ["", "", "", ""],
], widths_cm=[3.0, 2.2, 4.8, 6.8], header_bg=SKY_BG, header_color=SKY)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

section_chip(doc, "icon_share_final.png", "2. 가장 공감된 해설 카드", 3)
callout_box(doc, ["문장:", "", "", "이유:", "", ""], PURPLE, PURPLE_BG)

section_chip(doc, "icon_goal_final.png", "3. 게임과 우리 반 연결하기", 0)
callout_box(doc, ["게임 속 마을을 덮고 있던 '안개'처럼, 우리 반을 덮고 있는 안개는 무엇인가요?", "", ""], CORAL, CORAL_BG)
callout_box(doc, ["그 안개를 걷어내기 위해, 우리 반에 옮기고 싶은 약속 한 가지는?", "", ""], MINT, MINT_BG)
callout_box(doc, ["오늘 내 마음을 한 단어로 표현하면?  (                    )  — 그 이유는?", ""], GOLD, GOLD_BG)

section_chip(doc, "icon_badge_final.png", "4. 나의 탐험 기록", 1)
styled_grid(doc, [
    ["해결한 미션", "빛장벽", "마음 조각 💛", "우리 반 마을 밝기"],
    ["        / 16", "        / 144", "        / 24", "            %"],
], widths_cm=[4.2, 4.2, 4.2, 4.4], header_bg=MINT_BG, header_color=MINT)
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# 자기평가 별점 박스
star_table = doc.add_table(rows=2, cols=2)
no_table_borders(star_table)
set_col_widths(star_table, [8.5, 8.5])
labels = ["경청하며 참여했나요?", "친구의 이야기에 공감했나요?"]
for i, label in enumerate(labels):
    cell = star_table.rows[0].cells[i]
    set_cell_shading(cell, GOLD_BG)
    set_cell_borders(cell, color="E3C158", sz=6)
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = GOLD
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("☆ ☆ ☆ ☆ ☆")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0xE3, 0xB0, 0x1A)
prevent_row_split(star_table)
doc.add_paragraph().paragraph_format.space_after = Pt(10)

# 교사 확인 박스 (실제 도장/서명용 빈 칸)
confirm = doc.add_table(rows=1, cols=1)
no_table_borders(confirm)
ccell = confirm.rows[0].cells[0]
set_cell_shading(ccell, "FFFFFF")
set_cell_borders(ccell, color="C9BEB0", sz=6)
set_cell_margins(ccell, top=260, bottom=260, left=200, right=200)
cp = ccell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run("선생님 확인 · 도장")
cr.font.size = Pt(10)
cr.font.color.rgb = GRAY
prevent_row_split(confirm)

divider(doc)
doc.save(os.path.join(ROOT, "공동체빌더스_학생_활동지.docx"))
print("활동지 저장 완료")
